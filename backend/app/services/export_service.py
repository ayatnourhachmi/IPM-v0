"""Document export builders for delivery recommendations (PDF and DOCX)."""

from __future__ import annotations

import re
import shutil
import subprocess
from datetime import datetime
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from xml.sax.saxutils import escape

from app.core.recommendation_limits import (
    MAX_KPIS_RECOMMENDATIONS,
    MAX_ORGANIZATIONAL_RECOMMENDATIONS,
    MAX_TECHNICAL_RECOMMENDATIONS,
)

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.lib.utils import ImageReader
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


def _service_root() -> Path:
    return Path(__file__).resolve().parent.parent


def resolve_dxc_logo_path() -> Path | None:
    """Prefer bundled asset, then repo frontend landing asset (local dev)."""

    svc = Path(__file__).resolve()
    candidates = [
        _service_root() / "assets" / "dxc_logo.png",
        svc.parents[3] / "frontend" / "public" / "landing-files" / "DXC-Logo-2025.png",
    ]
    for p in candidates:
        if p.is_file():
            return p
    return None


def _fallback_logo_png_bytes() -> BytesIO:
    """Raster header when no official PNG is on disk."""
    from PIL import Image, ImageDraw, ImageFont

    w, h = 720, 112
    img = Image.new("RGB", (w, h), (8, 10, 14))
    draw = ImageDraw.Draw(img)
    for x in range(w):
        r = int(32 + min(140, x // 4))
        g = int(80 + min(120, x // 6))
        b = int(200 - min(140, x // 5))
        draw.line([(x, h - 10), (x, h - 4)], fill=(min(r, 255), min(g, 255), min(b, 255)))

    font_paths = (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/opentype/noto/NotoSans-Bold.ttf",
        "C:/Windows/Fonts/segoeui.ttf",
    )
    sub_paths = (
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/opentype/noto/NotoSans-Regular.ttf",
        "C:/Windows/Fonts/segoeui.ttf",
    )
    font = None
    for fp in font_paths:
        try:
            font = ImageFont.truetype(fp, 34)
            break
        except OSError:
            continue
    if font is None:
        font = ImageFont.load_default()
    sub = None
    for fp in sub_paths:
        try:
            sub = ImageFont.truetype(fp, 15)
            break
        except OSError:
            continue
    if sub is None:
        sub = font

    draw.text((28, 20), "DXC Technology", fill=(250, 250, 252), font=font)
    draw.text((28, 64), "Innovation Process Model · Delivery recommendations", fill=(180, 187, 196), font=sub)

    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def _logo_flowable() -> RLImage:
    from PIL import Image as PILImage

    target_w = 2.65 * inch
    pth = resolve_dxc_logo_path()
    if pth is not None:
        im = PILImage.open(pth)
        tw, th = im.size
        h = target_w * th / tw
        return RLImage(str(pth), width=target_w, height=h)
    buf = _fallback_logo_png_bytes()
    im = PILImage.open(buf)
    buf.seek(0)
    tw, th = im.size
    h = target_w * th / tw
    return RLImage(ImageReader(buf), width=target_w, height=h)


def _organizational_line(item: object) -> str:
    if isinstance(item, dict):
        role = str(item.get("role", "") or "").strip()
        action = str(item.get("action", "") or "").strip()
        if role and action:
            return f"{role} — {action}"
        return action or role or "Not specified"
    return str(item)


def _plain_summary(pitch: str, n_solutions: int, n_recs: int) -> str:
    p = (pitch or "").strip()
    intro = (
        "This document summarises delivery recommendations from your innovation qualification process. "
        "It brings together technical themes, organisational alignment, and measurable outcomes for the "
        "selected solutions."
    )
    scope = (
        f"{n_recs} recommendation section(s) address {n_solutions} solution(s) included in this export. "
        "Governance, prioritisation, and commitment remain with your portfolio and executive teams."
    )
    if len(p) > 420:
        p = p[:417] + "…"
    return f"{intro} {scope} <b>Need snapshot:</b> {p or 'Not specified.'}"


_META_TABLE_CAPTION = (
    "<i>Left: field label. Right: value for this export (reference and date).</i>"
)

_DELIVERY_TABLE_CAPTION = (
    "<b>Reading this table:</b> <b>Solution</b> — name of the recommended offering. "
    "<b>Relevance</b> — estimated semantic fit between your stated need and the solution (0–100%). "
    "<b>Overall score</b> — composite evaluation score from the qualification step (scale used in the workshop, typically 1–5)."
)

_KPI_TABLE_CAPTION = (
    "<b>Reading this table:</b> <b>KPI</b> — outcome or indicator to track. "
    "<b>Target</b> — level or threshold to reach within the agreed horizon. "
    "<b>How we measure</b> — evidence, metric source, or governance review where success is verified."
)


def _word_count(blob: str) -> int:
    return len(re.findall(r"\w+", blob or "", flags=re.UNICODE))


def _docx_set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()  # noqa: SLF001 - python-docx has no public API for shading.
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _docx_set_cell_text(cell, text: object, *, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text if text is not None else "—"))
    run.bold = bold
    run.font.size = Pt(9)
    if color is not None:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def _docx_style_table(table, *, header_fill: str = "1E3A8A", header_text: RGBColor | None = None) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if header_text is None:
        header_text = RGBColor(0xFF, 0xFF, 0xFF)
    for cell in table.rows[0].cells:
        _docx_set_cell_shading(cell, header_fill)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = header_text
    for row_index, row in enumerate(table.rows[1:], start=1):
        if row_index % 2 == 0:
            for cell in row.cells:
                _docx_set_cell_shading(cell, "F8FAFC")


def _docx_caption(document: Document, text: str) -> None:
    p = document.add_paragraph(text)
    if p.runs:
        p.runs[0].font.italic = True
        p.runs[0].font.size = Pt(8.5)
        p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)


def _docx_add_key_value_table(document: Document, rows: list[tuple[str, object]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows:
        cells = table.add_row().cells
        _docx_set_cell_text(cells[0], label, bold=True, color=RGBColor(0x33, 0x41, 0x55))
        _docx_set_cell_text(cells[1], value, color=RGBColor(0x1F, 0x29, 0x37))
        _docx_set_cell_shading(cells[0], "F1F5F9")


def _docx_add_bullets(document: Document, values: list[object], *, empty_text: str = "None listed.") -> None:
    if values:
        for item in values:
            document.add_paragraph(str(item), style="List Bullet")
    else:
        p = document.add_paragraph(empty_text)
        p.runs[0].italic = True


def _delivery_lookup(delivery_solutions: list[dict]) -> dict[str, dict]:
    return {str(solution.get("id") or ""): solution for solution in delivery_solutions}


def _find_soffice() -> str | None:
    for name in ("soffice", "libreoffice"):
        exe = shutil.which(name)
        if exe:
            return exe
    for path in (
        "C:/Program Files/LibreOffice/program/soffice.exe",
        "C:/Program Files (x86)/LibreOffice/program/soffice.exe",
    ):
        if Path(path).is_file():
            return path
    return None


def _convert_docx_bytes_to_pdf(docx_bytes: bytes, *, filename_stem: str) -> bytes:
    soffice = _find_soffice()
    if soffice is None:
        raise RuntimeError(
            "LibreOffice/soffice is required to export PDF from the canonical DOCX document."
        )

    safe_stem = re.sub(r"[^A-Za-z0-9_.-]+", "-", filename_stem).strip(".-") or "ipm-report"
    with TemporaryDirectory(prefix="ipm-export-") as tmp:
        tmp_path = Path(tmp)
        docx_path = tmp_path / f"{safe_stem}.docx"
        pdf_path = tmp_path / f"{safe_stem}.pdf"
        docx_path.write_bytes(docx_bytes)

        result = subprocess.run(
            [
                soffice,
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                "--convert-to",
                "pdf",
                "--outdir",
                str(tmp_path),
                str(docx_path),
            ],
            cwd=tmp_path,
            capture_output=True,
            text=True,
            timeout=90,
            check=False,
        )
        if result.returncode != 0 or not pdf_path.is_file():
            detail = (result.stderr or result.stdout or "no converter output").strip()
            raise RuntimeError(f"DOCX to PDF conversion failed: {detail}")
        return pdf_path.read_bytes()


def _pdf_table_paragraph(text: object, style: ParagraphStyle) -> Paragraph:
    """ReportLab Tables only wrap Flowables such as Paragraph; plain strings overlap."""
    s = str(text if text is not None else "").strip()
    return Paragraph(escape(s or "—"), style)


def _pdf_draw_footer(canvas, doc) -> None:  # noqa: ARG001
    canvas.saveState()
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.HexColor("#6B7280"))
    w, _ = A4
    canvas.drawCentredString(
        w / 2,
        34,
        "DXC Technology · Innovation Process Model · For client / internal use under engagement terms",
    )
    n = canvas.getPageNumber()
    canvas.drawRightString(w - 42, 34, f"Page {n}")
    canvas.restoreState()


def build_pdf_report(
    need_id: str,
    pitch: str,
    recommendations: list[dict],
    delivery_solutions: list[dict],
) -> bytes:
    docx_bytes = build_docx_report(
        need_id=need_id,
        pitch=pitch,
        recommendations=recommendations,
        delivery_solutions=delivery_solutions,
    )
    return _convert_docx_bytes_to_pdf(docx_bytes, filename_stem=f"{need_id}-recommendations")

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=48,
        rightMargin=48,
        topMargin=48,
        bottomMargin=56,
        title=f"IPM Recommendations - {need_id}",
        author="IPM / DXC",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#0F172A"),
        spaceAfter=6,
    )
    label_style = ParagraphStyle(
        "Label",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#374151"),
    )
    subtle_style = ParagraphStyle(
        "Subtle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#6B7280"),
        spaceAfter=14,
    )
    section_style = ParagraphStyle(
        "Section",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=12.5,
        leading=16,
        textColor=colors.HexColor("#111827"),
        spaceBefore=14,
        spaceAfter=8,
    )
    subhead_style = ParagraphStyle(
        "Subhead",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#1E3A5F"),
        spaceBefore=10,
        spaceAfter=6,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=14.5,
        textColor=colors.HexColor("#1F2937"),
        spaceAfter=4,
    )
    small_style = ParagraphStyle(
        "Small",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#64748B"),
        spaceAfter=8,
    )
    sol_tbl_hdr_para = ParagraphStyle(
        "SolTblHdrPara",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9,
        leading=11,
        textColor=colors.white,
    )
    sol_tbl_cell_para = ParagraphStyle(
        "SolTblCellPara",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=11.5,
        textColor=colors.HexColor("#1F2937"),
    )
    kpi_tbl_hdr_para = ParagraphStyle(
        "KpiTblHdrPara",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9,
        leading=11,
        textColor=colors.HexColor("#1E1B4B"),
    )
    kpi_tbl_cell_para = ParagraphStyle(
        "KpiTblCellPara",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=11.5,
        textColor=colors.HexColor("#1F2937"),
    )

    story: list = []
    story.append(_logo_flowable())
    story.append(Spacer(1, 10))
    story.append(
        Paragraph(
            "Innovation Process Model<br/><font color='#1E40AF'>PoC Preparation &amp; Delivery Recommendations</font>",
            title_style,
        )
    )
    story.append(Paragraph("<b><font color='#FF8A2A'>DXC Technology · Confidential</font></b>", label_style))
    story.append(Spacer(1, 8))

    gen_ts = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
    meta_rows = [
        ["Business Need Ref.", need_id],
        ["Document Version", "v1.0 — Draft"],
        ["Date", gen_ts],
        ["Document Owner", "DXC Engagement Lead"],
        ["Classification", "For client / internal use under engagement terms"],
    ]
    meta_tbl = Table(meta_rows, colWidths=[150, 340])
    meta_tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F1F5F9")),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#334155")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#E2E8F0")),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(meta_tbl)
    story.append(
        Paragraph(
            "Document control uses generated IPM data; client names, budgets, and signatures are intentionally left out until confirmed.",
            small_style,
        )
    )
    story.append(Spacer(1, 14))

    n_sol = len(delivery_solutions)
    summary = _plain_summary(pitch, n_sol, len(recommendations))
    story.append(Paragraph("1. Executive Summary", section_style))
    story.append(Paragraph(summary, body_style))

    story.append(Paragraph("Why DXC", subhead_style))
    for item in [
        "Structured Innovation Process Model (IPM) to guide need-to-production journeys.",
        "Enterprise delivery experience across cloud, data, AI, ERP, CRM, and service management ecosystems.",
        "Ability to convert qualification evidence into technical, organizational, and KPI-driven delivery actions.",
        "Governance approach designed for client / DXC alignment before production commitment.",
    ]:
        story.append(Paragraph(f"• {escape(item)}", body_style))

    story.append(Paragraph("2. Context & Framing", section_style))
    context_rows = [
        ["Business Need Ref.", need_id],
        ["Need Statement", pitch or "Not specified"],
        ["Recommendation Blocks", str(len(recommendations))],
        ["Solutions in Export", str(len(delivery_solutions))],
    ]
    context_tbl = Table(
        [[_pdf_table_paragraph(c, sol_tbl_cell_para) for c in row] for row in context_rows],
        colWidths=[145, 315],
    )
    context_tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F1F5F9")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(context_tbl)
    story.append(
        Paragraph(
            f"<i>Approx. {_word_count(pitch)} words in the need statement above.</i>",
            small_style,
        )
    )
    story.append(Spacer(1, 10))

    story.append(Paragraph("2.2 Scope Definition", subhead_style))
    story.append(
        Paragraph(
            "<b>In scope:</b> validate the selected solution recommendation bundle(s), confirm implementation prerequisites, "
            "define measurable KPIs, and prepare a PoC-oriented delivery path.",
            body_style,
        )
    )
    story.append(
        Paragraph(
            "<b>Out of scope:</b> final commercial terms, production architecture approval, client-specific budget, "
            "and named resource commitments.",
            body_style,
        )
    )

    if delivery_solutions:
        story.append(Paragraph("3. Selected Solution Portfolio", section_style))
        table_data = [
            [
                _pdf_table_paragraph("Solution", sol_tbl_hdr_para),
                _pdf_table_paragraph("Discovery Relevance", sol_tbl_hdr_para),
                _pdf_table_paragraph("IVI Score", sol_tbl_hdr_para),
            ]
        ]
        for solution in delivery_solutions:
            table_data.append(
                [
                    _pdf_table_paragraph(solution.get("name", "Unknown"), sol_tbl_cell_para),
                    _pdf_table_paragraph(f"{solution.get('relevance', 0)}%", sol_tbl_cell_para),
                    _pdf_table_paragraph(f"{float(solution.get('overall', 0)):.2f}", sol_tbl_cell_para),
                ]
            )
        sol_tbl = Table(table_data, colWidths=[280, 90, 90])
        sol_tbl.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A8A")),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(sol_tbl)
        story.append(
            Paragraph(
                "Discovery relevance is semantic fit to the need. IVI is the weighted qualification score used to prioritize candidates.",
                small_style,
            )
        )
        story.append(Spacer(1, 12))

    delivery_by_id = _delivery_lookup(delivery_solutions)
    for index, rec in enumerate(recommendations, start=1):
        story.append(PageBreak())

        title_line = rec.get("solution_name", "Solution")
        solution = delivery_by_id.get(str(rec.get("solution_id") or ""), {})
        story.append(Paragraph(f"4.{index} {escape(str(title_line))}", section_style))
        overview_rows = [
            ["Solution Name", title_line],
            ["Relevance Score", f"{solution.get('relevance', 0)}%" if solution else "Not specified"],
            ["IVI Score", f"{float(solution.get('overall', 0)):.2f}" if solution else "Not specified"],
            ["Recommendation Mode", str(rec.get("mode") or "STANDARD").upper()],
        ]
        overview_tbl = Table(
            [[_pdf_table_paragraph(c, sol_tbl_cell_para) for c in row] for row in overview_rows],
            colWidths=[145, 315],
        )
        overview_tbl.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F1F5F9")),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        story.append(overview_tbl)
        story.append(Spacer(1, 8))

        mode = str(rec.get("mode") or "STANDARD").upper()
        if mode == "PREREQUIS":
            story.append(
                Paragraph(
                    "<b>Recommendation mode: PREREQUIS</b> — solution fit is assessed as limited; "
                    "items below focus on readiness, prerequisites, and proofs—not a committed go-live plan.",
                    body_style,
                )
            )
            story.append(Spacer(1, 6))
        else:
            story.append(
                Paragraph(
                    "<b>Recommendation mode: STANDARD</b> — delivery-oriented guidance aligned with the "
                    "selected solution and your stated need.",
                    small_style,
                )
            )

        story.append(Paragraph("4.2 PoC Hypothesis", subhead_style))
        story.append(
            Paragraph(
                f"If {escape(str(title_line))} is applied to the stated business need, DXC expects measurable delivery value "
                "to be evidenced through the KPI targets and recommendation actions below.",
                body_style,
            )
        )

        story.append(Paragraph("4.3 Success Criteria", subhead_style))
        kpis = rec.get("kpis", [])[:MAX_KPIS_RECOMMENDATIONS]
        if kpis:
            for kpi in kpis:
                story.append(
                    Paragraph(
                        f"• {escape(str(kpi.get('name', 'KPI')))}: {escape(str(kpi.get('target', 'target to be confirmed')))}",
                        body_style,
                    )
                )
        else:
            story.append(Paragraph("Success criteria will be confirmed during PoC planning.", small_style))

        story.append(Paragraph("5. Technical Recommendations", section_style))
        story.append(
            Paragraph(
                "The recommendations below are generated from qualification scoring and gap-analysis evidence. "
                "They must be validated against client architecture before commitment.",
                body_style,
            )
        )
        tech = rec.get("technical_recommendations", [])[:MAX_TECHNICAL_RECOMMENDATIONS]
        for i, item in enumerate(tech, start=1):
            story.append(Paragraph(f"{i}. {escape(str(item))}", body_style))
        if not tech:
            story.append(Paragraph("— None listed —", small_style))
        story.append(Spacer(1, 8))

        story.append(Paragraph("6. Organizational Recommendations", section_style))
        story.append(Paragraph("6.1 Team & Roles", subhead_style))
        org = rec.get("organizational_recommendations", [])[:MAX_ORGANIZATIONAL_RECOMMENDATIONS]
        if org:
            org_rows = [
                [
                    _pdf_table_paragraph("Role", sol_tbl_hdr_para),
                    _pdf_table_paragraph("Recommended action / accountability", sol_tbl_hdr_para),
                ]
            ]
            for item in org:
                if isinstance(item, dict):
                    org_rows.append(
                        [
                            _pdf_table_paragraph(item.get("role", "Owner to confirm"), sol_tbl_cell_para),
                            _pdf_table_paragraph(item.get("action", "Action to confirm"), sol_tbl_cell_para),
                        ]
                    )
                else:
                    org_rows.append(
                        [
                            _pdf_table_paragraph("Owner to confirm", sol_tbl_cell_para),
                            _pdf_table_paragraph(item, sol_tbl_cell_para),
                        ]
                    )
            org_tbl = Table(org_rows, colWidths=[140, 320])
            org_tbl.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A8A")),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#CBD5E1")),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 8),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                        ("TOPPADDING", (0, 0), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.append(org_tbl)
        if not org:
            story.append(Paragraph("— None listed —", small_style))
        story.append(Spacer(1, 8))

        story.append(Paragraph("6.2 Governance & Steering", subhead_style))
        for item in [
            "Confirm a joint client / DXC owner for scope, acceptance testing, and decision rights.",
            "Review recommendation progress in the normal programme or portfolio cadence.",
            "Escalate architecture, data, and compliance blockers before committing to production scale.",
        ]:
            story.append(Paragraph(f"• {escape(item)}", body_style))

        story.append(Paragraph("7. KPI Definitions & Measurement", section_style))
        if kpis:
            kpi_header = [
                [
                    _pdf_table_paragraph("KPI", kpi_tbl_hdr_para),
                    _pdf_table_paragraph("Target", kpi_tbl_hdr_para),
                    _pdf_table_paragraph("How we measure", kpi_tbl_hdr_para),
                ]
            ]
            kpi_rows = []
            for kpi in kpis:
                kpi_rows.append(
                    [
                        _pdf_table_paragraph(kpi.get("name", "KPI"), kpi_tbl_cell_para),
                        _pdf_table_paragraph(kpi.get("target", "—"), kpi_tbl_cell_para),
                        _pdf_table_paragraph(kpi.get("measurement_criteria", "—"), kpi_tbl_cell_para),
                    ]
                )
            # Extra width on Target avoids wrapped lines colliding with the next column (plain strings overflow).
            kt = Table(kpi_header + kpi_rows, colWidths=[100, 182, 164])
            kt.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEF2FF")),
                        ("GRID", (0, 0), (-1, -1), 0.2, colors.HexColor("#C7D2FE")),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(kt)
            story.append(Paragraph("KPI baselines should be established and agreed before PoC kick-off.", small_style))
        else:
            story.append(Paragraph("— None listed —", small_style))

        story.append(Paragraph("8. Risk & Compliance", section_style))
        risk_rows = [
            [
                _pdf_table_paragraph("Risk", kpi_tbl_hdr_para),
                _pdf_table_paragraph("Category", kpi_tbl_hdr_para),
                _pdf_table_paragraph("Impact", kpi_tbl_hdr_para),
                _pdf_table_paragraph("Mitigation", kpi_tbl_hdr_para),
            ]
        ]
        for risk, category, impact, mitigation in [
            ("Data access delays", "Organizational", "High", "Agree access protocol and owners before PoC kick-off."),
            ("Data quality below threshold", "Technical", "High", "Run data profiling and define minimum quality criteria."),
            ("Scope creep", "Delivery", "Medium", "Use written change control for new scenarios or integrations."),
            ("Compliance blocker", "Compliance", "Critical", "Validate privacy, security, and data residency constraints before data processing."),
        ]:
            risk_rows.append(
                [
                    _pdf_table_paragraph(risk, kpi_tbl_cell_para),
                    _pdf_table_paragraph(category, kpi_tbl_cell_para),
                    _pdf_table_paragraph(impact, kpi_tbl_cell_para),
                    _pdf_table_paragraph(mitigation, kpi_tbl_cell_para),
                ]
            )
        risk_tbl = Table(risk_rows, colWidths=[110, 95, 65, 190])
        risk_tbl.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EEF2FF")),
                    ("GRID", (0, 0), (-1, -1), 0.2, colors.HexColor("#C7D2FE")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        story.append(risk_tbl)

        story.append(Paragraph("9. Next Steps & Path to Production", section_style))
        story.append(
            Paragraph(
                "A successful PoC should trigger client / DXC sign-off on measured outcomes, architecture readiness, "
                "production investment, and commercial next-phase terms.",
                body_style,
            )
        )
        for item in [
            "Confirm PoC owner, technical lead, data owner, and steering cadence.",
            "Validate the KPI baseline and measurement source.",
            "Review architecture, integration, security, and compliance prerequisites.",
            "Decide whether the candidate should move to production planning, rework, or stop.",
        ]:
            story.append(Paragraph(f"• {escape(item)}", body_style))

        story.append(Spacer(1, 12))

    doc.build(story, onFirstPage=_pdf_draw_footer, onLaterPages=_pdf_draw_footer)
    return buffer.getvalue()


def build_docx_report(
    need_id: str,
    pitch: str,
    recommendations: list[dict],
    delivery_solutions: list[dict],
) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    logo_path = resolve_dxc_logo_path()
    if logo_path is not None:
        p_logo = document.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_logo.add_run()
        run.add_picture(str(logo_path), width=Inches(2.6))
    else:
        buf = _fallback_logo_png_bytes()
        p_logo = document.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p_logo.add_run()
        run.add_picture(buf, width=Inches(2.6))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_title = title.add_run("Innovation Process Model\n")
    r_title.bold = True
    r_title.font.size = Pt(15)
    r_title.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    r_sub = title.add_run("PoC Preparation & Delivery Recommendations")
    r_sub.bold = True
    r_sub.font.size = Pt(22)
    r_sub.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    doc_note = document.add_paragraph()
    doc_note.add_run("DXC Technology · Confidential").bold = True
    doc_note.runs[0].font.size = Pt(9)
    doc_note.runs[0].font.color.rgb = RGBColor(0xFF, 0x8A, 0x2A)

    h1 = document.add_heading("1. Executive Summary", level=1)
    try:
        h1.runs[0].font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    except (IndexError, AttributeError):
        pass

    gen_ts = datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
    _docx_add_key_value_table(
        document,
        [
            ("Business Need Ref.", need_id),
            ("Document Version", "v1.0 — Draft"),
            ("Date", gen_ts),
            ("Document Owner", "DXC Engagement Lead"),
            ("Classification", "For client / internal use under engagement terms"),
        ],
    )
    _docx_caption(document, "Document control uses generated IPM data; client names, budgets, and signatures are intentionally left out until confirmed.")

    document.add_paragraph(
        _plain_summary(pitch, len(delivery_solutions), len(recommendations)).replace("<b>", "").replace("</b>", "")
    )

    document.add_heading("Why DXC", level=2)
    _docx_add_bullets(
        document,
        [
            "Structured Innovation Process Model (IPM) to guide need-to-production journeys.",
            "Enterprise delivery experience across cloud, data, AI, ERP, CRM, and service management ecosystems.",
            "Ability to convert qualification evidence into technical, organizational, and KPI-driven delivery actions.",
            "Governance approach designed for client / DXC alignment before production commitment.",
        ],
    )

    document.add_heading("2. Context & Framing", level=1)
    document.add_heading("2.1 Business Problem Statement", level=2)
    _docx_add_key_value_table(
        document,
        [
            ("Business Need Ref.", need_id),
            ("Need Statement", pitch or "Not specified"),
            ("Recommendation Blocks", len(recommendations)),
            ("Solutions in Export", len(delivery_solutions)),
        ],
    )
    _docx_caption(document, f"Approx. {_word_count(pitch)} words in the need statement above.")

    document.add_heading("2.2 Scope Definition", level=2)
    document.add_paragraph(
        "In scope: validate the selected solution recommendation bundle(s), confirm implementation prerequisites, "
        "define measurable KPIs, and prepare a PoC-oriented delivery path."
    )
    document.add_paragraph(
        "Out of scope: final commercial terms, production architecture approval, client-specific budget, and named resource commitments."
    )

    document.add_heading("3. Selected Solution Portfolio", level=1)
    if delivery_solutions:
        table = document.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Solution"
        hdr[1].text = "Discovery Relevance"
        hdr[2].text = "IVI Score"
        for solution in delivery_solutions:
            row = table.add_row().cells
            row[0].text = str(solution.get("name", "Unknown"))
            row[1].text = f"{solution.get('relevance', 0)}%"
            row[2].text = f"{float(solution.get('overall', 0)):.2f}"
        _docx_style_table(table)
        _docx_caption(document, "Discovery relevance is semantic fit to the need. IVI is the weighted qualification score used to prioritize candidates.")
    else:
        document.add_paragraph("No selected delivery solutions were provided.")

    delivery_by_id = _delivery_lookup(delivery_solutions)
    for index, rec in enumerate(recommendations, start=1):
        document.add_page_break()
        solution_id = str(rec.get("solution_id") or "")
        solution = delivery_by_id.get(solution_id, {})
        heading = f"4.{index} {rec.get('solution_name', 'Solution')}"
        document.add_heading(heading, level=1)

        document.add_heading("4.1 Solution Overview", level=2)
        _docx_add_key_value_table(
            document,
            [
                ("Solution Name", rec.get("solution_name", "Solution")),
                ("Relevance Score", f"{solution.get('relevance', 0)}%" if solution else "Not specified"),
                ("IVI Score", f"{float(solution.get('overall', 0)):.2f}" if solution else "Not specified"),
                ("Recommendation Mode", str(rec.get("mode") or "STANDARD").upper()),
            ],
        )

        mode = str(rec.get("mode") or "STANDARD").upper()
        p_mode = document.add_paragraph()
        if mode == "PREREQUIS":
            r0 = p_mode.add_run("Recommendation mode: PREREQUIS — ")
            r0.bold = True
            p_mode.add_run(
                "solution fit is assessed as limited; items below focus on readiness, prerequisites, and "
                "proofs—not a committed go-live plan."
            )
        else:
            r0 = p_mode.add_run("Recommendation mode: STANDARD — ")
            r0.bold = True
            p_mode.add_run(
                "delivery-oriented guidance aligned with the selected solution and your stated need."
            )

        document.add_heading("4.2 PoC Hypothesis", level=2)
        document.add_paragraph(
            f"If {rec.get('solution_name', 'the selected solution')} is applied to the stated business need, "
            "DXC expects measurable delivery value to be evidenced through the KPI targets and recommendation actions below."
        )

        document.add_heading("4.3 Success Criteria", level=2)
        kpis = rec.get("kpis", [])[:MAX_KPIS_RECOMMENDATIONS]
        if kpis:
            for kpi in kpis:
                document.add_paragraph(
                    f"{kpi.get('name', 'KPI')}: {kpi.get('target', 'target to be confirmed')}",
                    style="List Bullet",
                )
        else:
            document.add_paragraph("Success criteria will be confirmed during PoC planning.").runs[0].italic = True

        document.add_heading("5. Technical Recommendations", level=1)
        document.add_paragraph(
            "The recommendations below are generated from qualification scoring and gap-analysis evidence. "
            "They must be validated against client architecture before commitment."
        )
        tech = rec.get("technical_recommendations", [])[:MAX_TECHNICAL_RECOMMENDATIONS]
        _docx_add_bullets(document, tech)

        document.add_heading("6. Organizational Recommendations", level=1)
        document.add_heading("6.1 Team & Roles", level=2)
        org = rec.get("organizational_recommendations", [])[:MAX_ORGANIZATIONAL_RECOMMENDATIONS]
        if org:
            org_table = document.add_table(rows=1, cols=2)
            org_table.rows[0].cells[0].text = "Role"
            org_table.rows[0].cells[1].text = "Recommended action / accountability"
            for item in org:
                row = org_table.add_row().cells
                if isinstance(item, dict):
                    row[0].text = str(item.get("role") or "Owner to confirm")
                    row[1].text = str(item.get("action") or "Action to confirm")
                else:
                    row[0].text = "Owner to confirm"
                    row[1].text = str(item)
            _docx_style_table(org_table)
        if not org:
            document.add_paragraph("None listed.").runs[0].italic = True

        document.add_heading("6.2 Governance & Steering", level=2)
        _docx_add_bullets(
            document,
            [
                "Confirm a joint client / DXC owner for scope, acceptance testing, and decision rights.",
                "Review recommendation progress in the normal programme or portfolio cadence.",
                "Escalate architecture, data, and compliance blockers before committing to production scale.",
            ],
        )

        document.add_heading("7. KPI Definitions & Measurement", level=1)
        if kpis:
            t = document.add_table(rows=1, cols=3)
            h = t.rows[0].cells
            h[0].text = "KPI"
            h[1].text = "Target"
            h[2].text = "How we measure"
            for kpi in kpis:
                r = t.add_row().cells
                r[0].text = str(kpi.get("name", "KPI"))
                r[1].text = str(kpi.get("target", "—"))
                r[2].text = str(kpi.get("measurement_criteria", "—"))
            for row in t.rows:
                row.cells[0].width = Inches(1.15)
                row.cells[1].width = Inches(2.1)
                row.cells[2].width = Inches(2.2)
            _docx_style_table(t, header_fill="EEF2FF", header_text=RGBColor(0x1E, 0x1B, 0x4B))
            _docx_caption(document, "KPI baselines should be established and agreed before PoC kick-off.")
        else:
            document.add_paragraph("None listed.").runs[0].italic = True

        document.add_heading("8. Risk & Compliance", level=1)
        risk_table = document.add_table(rows=1, cols=4)
        risk_table.rows[0].cells[0].text = "Risk"
        risk_table.rows[0].cells[1].text = "Category"
        risk_table.rows[0].cells[2].text = "Impact"
        risk_table.rows[0].cells[3].text = "Mitigation"
        for risk, category, impact, mitigation in [
            ("Data access delays", "Organizational", "High", "Agree access protocol and owners before PoC kick-off."),
            ("Data quality below threshold", "Technical", "High", "Run data profiling and define minimum quality criteria."),
            ("Scope creep", "Delivery", "Medium", "Use written change control for new scenarios or integrations."),
            ("Compliance blocker", "Compliance", "Critical", "Validate privacy, security, and data residency constraints before data processing."),
        ]:
            row = risk_table.add_row().cells
            row[0].text = risk
            row[1].text = category
            row[2].text = impact
            row[3].text = mitigation
        _docx_style_table(risk_table)

        document.add_heading("9. Next Steps & Path to Production", level=1)
        document.add_paragraph(
            "A successful PoC should trigger client / DXC sign-off on measured outcomes, architecture readiness, "
            "production investment, and commercial next-phase terms."
        )
        _docx_add_bullets(
            document,
            [
                "Confirm PoC owner, technical lead, data owner, and steering cadence.",
                "Validate the KPI baseline and measurement source.",
                "Review architecture, integration, security, and compliance prerequisites.",
                "Decide whether the candidate should move to production planning, rework, or stop.",
            ],
        )

    try:
        section.footer.paragraphs[0].text = (
            "DXC Technology · Innovation Process Model · For client / internal use under engagement terms"
        )
    except Exception:
        pass

    output = BytesIO()
    document.save(output)
    return output.getvalue()
